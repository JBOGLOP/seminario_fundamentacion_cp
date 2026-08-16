"""
Genera el Programador del Seminario de fundamentación en cuidados paliativos
rellenando la plantilla institucional «4. Formato. Programador.docx».

Se rellena la plantilla real, no se recrea: así el documento conserva los
estilos, márgenes y encabezados de la Vicerrectoría Académica.

Fuentes:
  · Contenido: recursos/originales/Seminario de fundamentación...docx
  · Calendario: recursos/originales/HORARIOS 2026 - II.pdf
  · Convención de diligenciamiento: Programador Contexto histórico 26-1.docx
"""
import copy, sys, pathlib
import docx
from docx.shared import Pt

RAIZ = pathlib.Path(__file__).resolve().parent.parent
PLANTILLA = RAIZ / 'recursos' / 'originales' / '4. Formato. Programador.docx'
SALIDA = RAIZ / 'docs' / 'Programador_2026-II.docx'

DOCENTE = 'Jorge Wilhem Bogoya López'

# ── Datos de identificación ───────────────────────────────────────────
# Las horas se toman del contenido programático oficial (192/48/144). El
# espacio es teórico, así que las teóricas igualan a las presenciales y no
# hay prácticas. Ver la nota al pie del documento.
DATOS = {
    1: ('Programa: \nMaestría en Cuidados Paliativos',
        'Asignatura:\nSeminario de fundamentación en cuidados paliativos'),
    2: ('Código:  \n32264001', 'Plan de estudios:\n3660'),
    3: ('Número de Créditos dentro del Plan de Estudios: 3',
        'Fecha de actualización:\nAgosto de 2026'),
    4: ('Horas Totales', '192'),
    5: ('Horas presenciales', '48'),
    6: ('Horas de trabajo independiente', '144'),
    7: ('Horas teóricas', '48'),
    8: ('Horas prácticas', '0'),
}

# ── Contenidos ────────────────────────────────────────────────────────
# Un renglón por bloque horario, como en el programador de Contexto
# Histórico 26-1. Franja del seminario: viernes de 8:00 a 11:00.
# (semana, fecha, bloque, tema, responsable)
FILAS = [
    # ── Sesión 1 · Unidad 1 · Tema 1 ──────────────────────────────────
    ('4', '21/08/2026', '8:00-9:00 am',
     'Encuentro presencial\n'
     'Presentación del espacio académico\n'
     'Unidad 1. Contexto Histórico-filosófico y lineamientos sociopolíticos de los '
     'cuidados paliativos\n'
     'Tema 1. Antecedentes históricos de los cuidados paliativos', DOCENTE),
    ('', '21/08/2026', '9:00-10:00 am',
     'Tema 1. Aspectos históricos, desarrollo y tendencias de los cuidados paliativos', ''),
    ('', '21/08/2026', '10:00-11:00 am',
     'Tema 1. Fundamentos filosóficos de los cuidados paliativos', ''),

    # ── Sesión 2 · Unidad 1 · Tema 2 ──────────────────────────────────
    ('6', '04/09/2026', '8:00-9:00 am',
     'Encuentro sincrónico\n'
     'Tema 2. Lineamientos y políticas mundiales de los cuidados paliativos '
     '(OMS, Banco Mundial, BID, OCDE, IAHPC, SECPAL)', DOCENTE),
    ('', '04/09/2026', '9:00-10:00 am',
     'Tema 2. Lineamientos y políticas en cuidados paliativos a nivel latinoamericano '
     '(OPS, CEPAL, ALCP)', ''),
    ('', '04/09/2026', '10:00-11:00 am',
     'Tema 2. Lineamientos y políticas en cuidados paliativos en Colombia '
     '(PAIS, Modelo Preventivo y resolutivo, RIAS)', ''),

    # ── Sesión 3 · Unidad 1 · Tema 3 · cierre primer corte ────────────
    ('8', '18/09/2026', '8:00-9:00 am',
     'Encuentro presencial\n'
     'Tema 3. Marco legal de los cuidados paliativos\n'
     'Legislación Internacional y Latinoamérica en cuidados paliativos', DOCENTE),
    ('', '18/09/2026', '9:00-10:00 am',
     'Tema 3. Legislación Nacional alrededor de los cuidados paliativos\n'
     'Enfoque de derechos humanos en cuidados paliativos', ''),
    ('', '18/09/2026', '10:00-11:00 am',
     'Tema 3. Barreras de acceso en el contexto de cuidado paliativo\n'
     'Entrega Primer Producto 35 %', ''),

    # ── Sesión 4 · Unidad 2 · Temas 1 y 2 ─────────────────────────────
    ('10', '02/10/2026', '8:00-9:00 am',
     'Encuentro sincrónico\n'
     'Unidad 2. Modelos de atención en cuidados paliativos desde la investigación\n'
     'Tema 1. Modelos tradicionales vs. alternativos de atención en unidades de '
     'cuidados paliativos', DOCENTE),
    ('', '02/10/2026', '9:00-10:00 am',
     'Tema 1. Lectura crítica de la evidencia que sustenta cada modelo de atención', ''),
    ('', '02/10/2026', '10:00-11:00 am',
     'Tema 2. Incorporación de telesalud y tecnologías digitales', ''),

    # ── Sesión 5 · Unidad 2 · Tema 3 · cierre segundo corte ───────────
    ('12', '16/10/2026', '8:00-9:00 am',
     'Encuentro presencial\n'
     'Tema 3. Enfoques diferenciales: servicio domiciliario', DOCENTE),
    ('', '16/10/2026', '9:00-10:00 am',
     'Tema 3. Enfoques diferenciales: servicio institucional y comunitario', ''),
    ('', '16/10/2026', '10:00-11:00 am',
     'Tema 3. Cierre de la unidad y discusión de casos\n'
     'Entrega Segundo Producto 35 %', ''),

    # ── Sesión 6 · Unidad 3 · Tema 1 ──────────────────────────────────
    ('14', '30/10/2026', '8:00-9:00 am',
     'Encuentro sincrónico\n'
     'Unidad 3. Aspectos bioéticos y experiencias de manejo integral en cuidado '
     'paliativo desde la investigación\n'
     'Tema 1. Modelos de atención bioética en cuidado paliativo', DOCENTE),
    ('', '30/10/2026', '9:00-10:00 am',
     'Tema 1. Aspectos relevantes de los dilemas éticos en cuidado paliativo', ''),
    ('', '30/10/2026', '10:00-11:00 am',
     'Tema 1. Planificación anticipada de la atención y decisiones de cuidado\n'
     'Aspectos legales y directrices anticipadas', ''),

    # ── Encuentro institucional, ajeno a este espacio académico ───────
    ('16', '13/11/2026', '',
     'Encuentro de investigación — sustentaciones de tesis\n'
     '(Solo para estudiantes de Tesis II. No corresponde a este espacio académico)', ''),

    # ── Sesión 7 · Unidad 3 · Tema 2 · cierre ─────────────────────────
    ('17', '20/11/2026', '8:00-9:00 am',
     'Encuentro presencial\n'
     'Tema 2. Estrategias y experiencias de manejo integral en cuidado paliativo\n'
     'Funciones del equipo interprofesional y liderazgo colaborativo', DOCENTE),
    ('', '20/11/2026', '9:00-10:00 am',
     'Tema 2. Coordinación de recursos comunitarios y redes de apoyo\n'
     'Programas de manejo exitosos en cuidado paliativo', ''),
    ('', '20/11/2026', '10:00-11:00 am',
     'Cierre del espacio académico\n'
     'Entrega Producto Final 30 %', ''),
]

NOTA = (
    'Nota. La programación sigue el calendario oficial de encuentros académicos 2026-II de la '
    'Maestría (comunicación de la Dirección) y la franja del seminario, viernes de 8:00 a 11:00. '
    'Sobre esa base el semestre ofrece siete sesiones. Quedan dos puntos por confirmar con la '
    'Dirección: (1) el encuentro del 8 de agosto de 2026 cae en sábado, único del semestre sin '
    'viernes, por lo que no se programó como sesión de este espacio académico; de confirmarse, '
    'las sesiones pasarían a ocho y la telesalud recuperaría tema propio. (2) Las 48 horas '
    'presenciales declaradas en el contenido programático corresponden a 21 horas de encuentro '
    'efectivas (7 × 3 h); se registran las declaradas por coherencia con el contenido programático '
    'aprobado.'
)


# ── Utilidades sobre la plantilla ─────────────────────────────────────
def escribir(celda, texto):
    """Escribe en la celda respetando el formato del primer run existente."""
    parrafos = celda.paragraphs
    base = parrafos[0]
    modelo = base.runs[0] if base.runs else None

    def clonar_formato(run):
        if modelo is None:
            return
        run.font.name = modelo.font.name
        run.font.size = modelo.font.size
        run.font.bold = modelo.font.bold
        run.font.italic = modelo.font.italic
        if modelo.font.color and modelo.font.color.rgb:
            run.font.color.rgb = modelo.font.color.rgb

    # Vaciar todos los párrafos salvo el primero
    for p in parrafos[1:]:
        p._element.getparent().remove(p._element)
    for r in list(base.runs):
        r._element.getparent().remove(r._element)

    lineas = texto.split('\n')
    run = base.add_run(lineas[0])
    clonar_formato(run)
    for linea in lineas[1:]:
        nuevo = copy.deepcopy(base._element)
        base._element.addnext(nuevo)
        base = docx.text.paragraph.Paragraph(nuevo, base._parent)
        for r in list(base.runs):
            r._element.getparent().remove(r._element)
        run = base.add_run(linea)
        clonar_formato(run)


def main():
    doc = docx.Document(str(PLANTILLA))
    tabla = doc.tables[0]

    # 1 · Datos de identificación
    for indice, (etiqueta, valor) in DATOS.items():
        fila = tabla.rows[indice]
        escribir(fila.cells[0], etiqueta)
        escribir(fila.cells[3], valor)

    # 2 · Contenidos. Las filas de contenido de la plantilla van de la 12
    #     en adelante; se reutilizan las que hagan falta y se borra el resto.
    primera = 12
    disponibles = len(tabla.rows) - primera
    if len(FILAS) > disponibles:
        print('ERROR: la plantilla tiene %d filas de contenido y se necesitan %d'
              % (disponibles, len(FILAS)))
        return 1

    for i, (semana, fecha, bloque, tema, responsable) in enumerate(FILAS):
        fila = tabla.rows[primera + i]
        escribir(fila.cells[0], semana)
        escribir(fila.cells[1], fecha + ('\n' + bloque if bloque else ''))
        escribir(fila.cells[2], tema)      # celdas 2 y 3 van combinadas
        escribir(fila.cells[4], responsable)

    # Borrar las filas sobrantes de la plantilla (PRIMER PARCIAL, ROTACIÓN,
    # EXAMEN FINAL y los renglones en blanco): este espacio académico evalúa
    # por productos de corte, no por parciales, y no tiene rotaciones.
    for fila in list(tabla.rows)[primera + len(FILAS):]:
        fila._tr.getparent().remove(fila._tr)

    # 3 · Nota al pie
    p = doc.add_paragraph()
    run = p.add_run(NOTA)
    run.font.size = Pt(8)
    run.font.italic = True

    doc.save(str(SALIDA))
    print('Escrito: %s' % SALIDA)
    print('  filas de contenido: %d' % len(FILAS))
    print('  filas totales en la tabla: %d' % len(tabla.rows))
    return 0


if __name__ == '__main__':
    sys.exit(main())
